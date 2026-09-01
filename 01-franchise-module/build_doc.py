# -*- coding: utf-8 -*-
import os, math
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = os.path.dirname(os.path.abspath(__file__))
IMG = os.path.join(HERE, "imgs")
os.makedirs(IMG, exist_ok=True)
DOCX = os.path.join(HERE, "Franchise_Module_Documentation.docx")

SS = 3
FREG = r"C:\Windows\Fonts\segoeui.ttf"
FBLD = r"C:\Windows\Fonts\segoeuib.ttf"
def F(sz, bold=True): return ImageFont.truetype(FBLD if bold else FREG, int(sz*SS))

COL = {
 'wh':(30,58,138),'mw':(124,58,237),'l1':(37,99,235),'l2':(8,145,178),
 'l3':(5,150,105),'u':(245,158,11),'admin':(219,39,119),'red':(220,38,38),
 'green':(5,150,105),'line':(100,116,139),'ink':(15,23,42),'bg':(248,250,252),
}
def hx(rgb): return rgb

def clip(c,w,h,p):
    dx=p[0]-c[0]; dy=p[1]-c[1]
    if dx==0 and dy==0: return c
    sx=(w/2)/abs(dx) if dx else 1e9
    sy=(h/2)/abs(dy) if dy else 1e9
    s=min(sx,sy)
    return (c[0]+dx*s, c[1]+dy*s)

def head(d,frm,to,color,size=11):
    ang=math.atan2(to[1]-frm[1],to[0]-frm[0]); s=size*SS
    p1=(to[0]-s*math.cos(ang-0.42),to[1]-s*math.sin(ang-0.42))
    p2=(to[0]-s*math.cos(ang+0.42),to[1]-s*math.sin(ang+0.42))
    d.polygon([to,p1,p2],fill=color)

def dline(d,p1,p2,color,w,dashed):
    if not dashed:
        d.line([p1,p2],fill=color,width=w); return
    dist=math.hypot(p2[0]-p1[0],p2[1]-p1[1]); n=max(1,int(dist/(9*SS)))
    for i in range(n):
        if i%2: continue
        a=(p1[0]+(p2[0]-p1[0])*i/n, p1[1]+(p2[1]-p1[1])*i/n)
        b=(p1[0]+(p2[0]-p1[0])*(i+1)/n, p1[1]+(p2[1]-p1[1])*(i+1)/n)
        d.line([a,b],fill=color,width=w)

def tc(d,cx,cy,lines,font,fill):
    lh=font.size+5*SS; y=cy-(len(lines)*lh)/2
    for ln in lines:
        bb=d.textbbox((0,0),ln,font=font); tw=bb[2]-bb[0]
        d.text((cx-tw/2,y),ln,font=font,fill=fill); y+=lh

def draw(name,W,H,nodes,edges,nw=150,nh=64):
    img=Image.new("RGB",(W*SS,H*SS),COL['bg']); d=ImageDraw.Draw(img)
    nf=F(13); lf=F(11)
    for e in edges:
        a,b=e[0],e[1]; label=e[2] if len(e)>2 else None
        dashed=e[3] if len(e)>3 else False
        ck=e[4] if len(e)>4 else 'line'; ecol=COL.get(ck,COL['line'])
        off=e[5] if len(e)>5 else (0,0)
        na,nb=nodes[a],nodes[b]
        c1=((na['x']+off[0])*SS,(na['y']+off[1])*SS)
        c2=((nb['x']+off[0])*SS,(nb['y']+off[1])*SS)
        wa=na.get('w',nw)*SS; ha=na.get('h',nh)*SS; wb=nb.get('w',nw)*SS; hb=nb.get('h',nh)*SS
        p1=clip(c1,wa,ha,c2); p2=clip(c2,wb,hb,c1)
        dline(d,p1,p2,ecol,3*SS if not dashed else 2*SS,dashed); head(d,p1,p2,ecol)
        if label:
            mx=(p1[0]+p2[0])/2; my=(p1[1]+p2[1])/2
            bb=d.textbbox((0,0),label,font=lf); tw=bb[2]-bb[0]; th=bb[3]-bb[1]; pad=5*SS
            d.rounded_rectangle([mx-tw/2-pad,my-th/2-pad,mx+tw/2+pad,my+th/2+pad],radius=6*SS,fill=(255,255,255),outline=(203,213,225),width=SS)
            d.text((mx-tw/2,my-th/2-2*SS),label,font=lf,fill=COL['ink'])
    for nid,n in nodes.items():
        w=n.get('w',nw)*SS; h=n.get('h',nh)*SS; x=n['x']*SS; y=n['y']*SS; col=COL[n['c']]
        d.rounded_rectangle([x-w/2,y-h/2,x+w/2,y+h/2],radius=14*SS,fill=col)
        tc(d,x,y,n['t'].split('\n'),nf,(255,255,255))
        if n.get('badge'):
            bf=F(11); bt=n['badge']; bb=d.textbbox((0,0),bt,font=bf); tw=bb[2]-bb[0]; th=bb[3]-bb[1]
            bx=x+w/2+9*SS; by=y; pad=6*SS
            d.rounded_rectangle([bx,by-th/2-pad,bx+tw+2*pad,by+th/2+pad],radius=7*SS,fill=(254,243,199),outline=(245,158,11),width=SS)
            d.text((bx+pad,by-th/2-2*SS),bt,font=bf,fill=(146,64,14))
    img=img.resize((W,H),Image.LANCZOS); p=os.path.join(IMG,name+".png"); img.save(p); return p

# ---------------- diagrams ----------------
L1='l1'; L2='l2'; L3='l3'
def N(x,y,t,c,**k): d={'x':x,'y':y,'t':t,'c':c}; d.update(k); return d

# 1 hierarchy full
draw("d_hier",900,580,{
 'WH':N(430,55,'WH\nWarehouse','wh'),
 'MW':N(430,170,'Mini-Warehouse','mw'),
 'S1':N(120,300,'S1\nLevel 1',L1),'S2':N(390,300,'S2\nLevel 1',L1),'S3':N(680,300,'S3\nLevel 1',L1),
 'S31':N(580,440,'S3.1\nLevel 2',L2),'S311':N(790,440,'S3.2\nLevel 3',L3),
},[('WH','MW','stock'),('MW','S1',),('MW','S2',),('MW','S3',),('S3','S31',),('S3','S311',)])

# 2 multiple mini-wh
draw("d_multimw",820,400,{
 'WH':N(410,55,'WH\nWarehouse','wh'),
 'MW1':N(230,180,'Mini-WH 1','mw'),'MW2':N(590,180,'Mini-WH 2','mw'),
 'A1':N(120,320,'S1\nLevel 1',L1),'A2':N(340,320,'S2\nLevel 1',L1),
 'B1':N(480,320,'S5\nLevel 1',L1),'B2':N(700,320,'S6\nLevel 1',L1),
},[('WH','MW1',),('WH','MW2',),('MW1','A1',),('MW1','A2',),('MW2','B1',),('MW2','B2',)])

# 3 sparse 4 cases
draw("d_spA",400,250,{'MW':N(200,55,'Mini-WH','mw'),'S2':N(200,180,'S2\nLevel 1',L1)},[('MW','S2',)])
draw("d_spB",400,300,{'MW':N(200,55,'Mini-WH','mw'),'S1':N(200,160,'S1\nLevel 1',L1),'S11':N(200,260,'S1.1\nLevel 2',L2)},[('MW','S1',),('S1','S11',)])
draw("d_spC",400,330,{'MW':N(200,55,'Mini-WH','mw'),'S3':N(200,160,'S3\nLevel 1',L1),'S31':N(105,275,'S3.1\nLevel 2',L2),'S311':N(295,275,'S3.2\nLevel 3',L3)},[('MW','S3',),('S3','S31',),('S3','S311',)])
draw("d_spD",400,300,{'MW':N(200,55,'Mini-WH','mw'),'S4':N(200,160,'S4\nLevel 1',L1),'S411':N(200,260,'S4.2\nLevel 3',L3)},[('MW','S4',),('S4','S411',)])

# 4 pincode rollup
draw("d_pin",900,420,{
 'S31':N(175,320,'S3.1 (Level 2)',L2,badge='serve: sirf p4',w=210),
 'S311':N(560,320,'S3.2 (Level 3)',L3,badge='serve: p1,p2,p3',w=210),
 'S3':N(367,120,'S3 (Level 1)','l1',badge='own p5,p6 + p1..p4',w=210),
},[('S31','S3','rollup up'),('S311','S3','rollup up')],nw=210)

# 5 user match
draw("d_user",930,620,{
 'WH':N(420,50,'WH','wh'),'MW':N(420,165,'Mini-WH','mw',badge='p1 OK'),
 'S1':N(115,300,'S1','l1',badge='p8  X'),'S2':N(380,300,'S2','l1',badge='p9  X'),'S3':N(650,300,'S3','l1',badge='p1 OK'),
 'S31':N(530,445,'S3.1',L2,badge='p4  X'),'S311':N(755,445,'S3.2',L3,badge='p1 OK'),
 'U2':N(360,570,'User U2\npincode p1','u'),
},[('WH','MW',),('MW','S1',),('MW','S2',),('MW','S3',),('S3','S31',),('S3','S311',),('S311','U2','serve p1')])

# 6 limit decision flow
draw("d_limit",720,700,{
 'O':N(190,60,'User order\nRs X','l1',w=170),
 'Q':N(190,200,'Home shop?\n(pincode match)','u',w=200),
 'OK':N(520,200,'UNLIMITED\nallow','green',w=180),
 'E':N(190,360,'Effective limit\nspecial OR Rs2000','l2',w=210),
 'L':N(190,510,'spent + X within\nlimit?','u',w=200),
 'OK2':N(520,510,'Allow\nspent += X','green',w=170),
 'B':N(190,640,'BLOCK\nrequest to admin','red',w=200),
},[('O','Q',),('Q','OK','MATCH'),('Q','E','no match'),('E','L',),('L','OK2','yes'),('L','B','no')])

# 7 limit request <-> admin
draw("d_req",640,250,{
 'ADMIN':N(170,125,'Admin','admin'),'U2':N(470,125,'User','u'),
},[('U2','ADMIN','request (amount + reason)',False,'admin',(0,-22)),
   ('ADMIN','U2','approve -> limit increase',False,'green',(0,22))])

# 8 commission setup — do alag raaste (L2 aur L3 dono L1 ke neeche)
draw("d_setup",760,400,{
 'WH':N(380,70,'WH','wh',w=140),'S3':N(380,200,'Level 1','l1',w=150),
 'S31':N(230,330,'Level 2',L2,w=150),'S311':N(530,330,'Level 3',L3,w=150),
},[('WH','S3','Rs800'),('S3','S31','Rs800'),('S3','S311','Rs800')],nw=150)

# 9 raasta 1 — user L2 se le
draw("d_cA",760,240,{
 'WH':N(110,120,'WH','wh',w=130),'S3':N(310,120,'Level 1\napna hissa','l1',w=150),
 'S31':N(520,120,'Level 2\napna hissa',L2,w=150),'U2':N(700,120,'User\npays 800','u',w=120),
},[('WH','S3','800'),('S3','S31','800'),('S31','U2','800')],nw=150)

# 10 raasta 2 — user L3 se le
draw("d_cB",760,240,{
 'WH':N(110,120,'WH','wh',w=130),'S3':N(310,120,'Level 1\napna hissa','l1',w=150),
 'S311':N(520,120,'Level 3\napna hissa',L3,w=150),'U2':N(700,120,'User\npays 800','u',w=120),
},[('WH','S3','800'),('S3','S311','800'),('S311','U2','800')],nw=150)

# 11 case C (L1 direct 50/50)
draw("d_cC",680,240,{
 'WH':N(130,120,'WH\n+50  (50%)','wh',w=180),'S3':N(390,120,'Level 1\n+50  (50%)','l1',w=180),'U2':N(600,120,'User\npays 800','u',w=140),
},[('WH','S3','800'),('S3','U2','800')],nw=170)

# 13 case D (L1 sells 900)
draw("d_cD",700,240,{
 'WH':N(130,120,'WH\n+50','wh',w=160),'S3':N(380,120,'Level 1\n+150','l1',w=180),'U2':N(600,120,'User\npays 900','u',w=150),
},[('WH','S3','800'),('S3','U2','becha 900')],nw=170)

# 14 case E (L3 sells 900)
draw("d_cE",780,240,{
 'WH':N(110,120,'WH','wh',w=130),'S3':N(310,120,'Level 1\napna hissa','l1',w=150),
 'S311':N(520,120,'Level 3\napna +100*',L3,w=150),'U2':N(710,120,'User\npays 900','u',w=120),
},[('WH','S3','800'),('S3','S311','800'),('S311','U2','becha 900')],nw=150)

# 15 commission split - har level ko apna
draw("d_cF",760,400,{
 'WH':N(380,70,'WH','wh',w=140),'S3':N(380,200,'Level 1\napna','l1',w=160),
 'S31':N(230,330,'Level 2\napna',L2,w=160),'S311':N(530,330,'Level 3\napna',L3,w=160),
},[('WH','S3','800'),('S3','S31','800'),('S3','S311','800')],nw=160)

# 16 mini-wh commission
draw("d_cMW",680,240,{
 'WH':N(130,120,'WH','wh',w=150),'MW':N(370,120,'Mini-WH\n+15','mw',w=170),'S3':N(590,120,'Level 1','l1',w=150),
},[('WH','MW','800'),('MW','S3','800')],nw=160)

# 16b mini-wh direct to user (50/50 WH-MW)
draw("d_cMWuser",700,240,{
 'WH':N(140,120,'WH\n+50','wh',w=160),'MW':N(390,120,'Mini-WH\n+15 +50 = 65','mw',w=200),'U2':N(610,120,'User\npays 800','u',w=150),
},[('WH','MW','800'),('MW','U2','800')],nw=170)

# 16c L1 direct user inside MW chain (WH-L1 50/50, MW keeps 15)
draw("d_cL1mw",820,240,{
 'WH':N(100,120,'WH\n+50','wh',w=140),'MW':N(310,120,'Mini-WH\n+15','mw',w=150),'S3':N(530,120,'Level 1\n+50','l1',w=150),'U2':N(730,120,'User\npays 800','u',w=130),
},[('WH','MW','800'),('MW','S3','800'),('S3','U2','800')],nw=150)

# 17 invoice flow
draw("d_inv",780,420,{
 'WH':N(390,70,'WH','wh',w=130),'S3':N(390,200,'Level 1','l1',w=140),
 'S31':N(210,330,'Level 2',L2,w=140),'S311':N(560,330,'Level 3',L3,w=140),
},[('WH','S3','Stock Transfer'),('S3','S31','Tax Invoice'),('S3','S311','Tax Invoice')],nw=140)

print("diagrams done")

# ================= DOCX =================
doc=Document()
sec=doc.sections[0]
sec.top_margin=Inches(0.7); sec.bottom_margin=Inches(0.7); sec.left_margin=Inches(0.8); sec.right_margin=Inches(0.8)
PW=Inches(6.9)

def shade(cell,fill):
    tcPr=cell._tc.get_or_add_tcPr(); sh=OxmlElement('w:shd')
    sh.set(qn('w:val'),'clear'); sh.set(qn('w:fill'),fill); tcPr.append(sh)

def set_cell(cell,text,bold=False,color='1F2937',size=10.5,align='center',fill=None):
    cell.text=''; p=cell.paragraphs[0]
    p.alignment={'center':WD_ALIGN_PARAGRAPH.CENTER,'left':WD_ALIGN_PARAGRAPH.LEFT}[align]
    r=p.add_run(text); r.bold=bold; r.font.size=Pt(size); r.font.name='Segoe UI'
    r.font.color.rgb=RGBColor.from_string(color)
    if fill: shade(cell,fill)

def H(text,size=15,color='1E3A8A',before=14,after=4):
    p=doc.add_paragraph(); pf=p.paragraph_format; pf.space_before=Pt(before); pf.space_after=Pt(after)
    r=p.add_run(text); r.bold=True; r.font.size=Pt(size); r.font.name='Segoe UI'
    r.font.color.rgb=RGBColor.from_string(color); return p

def para(runs,after=6,size=10.5,align='left'):
    p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(after)
    p.alignment={'left':WD_ALIGN_PARAGRAPH.LEFT,'center':WD_ALIGN_PARAGRAPH.CENTER}[align]
    if isinstance(runs,str): runs=[(runs,False)]
    for t,b in runs:
        r=p.add_run(t); r.bold=b; r.font.size=Pt(size); r.font.name='Segoe UI'
        r.font.color.rgb=RGBColor.from_string('374151' if not b else '1E3A8A')
    return p

def bullet(text,bolds=None):
    p=doc.add_paragraph(style='List Bullet'); p.paragraph_format.space_after=Pt(3)
    r=p.add_run(text); r.font.size=Pt(10.5); r.font.name='Segoe UI'; r.font.color.rgb=RGBColor.from_string('374151')
    return p

def callout(title,text,fill='EFF6FF',tcolor='1E40AF'):
    t=doc.add_table(rows=1,cols=1); t.style='Table Grid'
    c=t.cell(0,0); shade(c,fill); c.paragraphs[0].text=''
    p=c.paragraphs[0]; r=p.add_run(title+'  '); r.bold=True; r.font.size=Pt(10.5); r.font.name='Segoe UI'; r.font.color.rgb=RGBColor.from_string(tcolor)
    r2=p.add_run(text); r2.font.size=Pt(10.5); r2.font.name='Segoe UI'; r2.font.color.rgb=RGBColor.from_string('374151')
    doc.add_paragraph().paragraph_format.space_after=Pt(2)

def pic(name,width=6.6,cap=None):
    doc.add_picture(os.path.join(IMG,name+'.png'),width=Inches(width))
    doc.paragraphs[-1].alignment=WD_ALIGN_PARAGRAPH.CENTER
    if cap:
        p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after=Pt(8)
        r=p.add_run(cap); r.italic=True; r.font.size=Pt(9); r.font.color.rgb=RGBColor.from_string('64748B')

def table(headers,rows,hfill='1E40AF',widths=None):
    t=doc.add_table(rows=1,cols=len(headers)); t.style='Table Grid'; t.alignment=WD_ALIGN_PARAGRAPH.CENTER
    for i,h in enumerate(headers): set_cell(t.rows[0].cells[i],h,bold=True,color='FFFFFF',fill=hfill)
    for ri,row in enumerate(rows):
        cells=t.add_row().cells
        for i,v in enumerate(row):
            set_cell(cells[i],str(v),fill='F1F5F9' if ri%2 else None)
    doc.add_paragraph().paragraph_format.space_after=Pt(2)
    return t

# ---------- COVER ----------
for _ in range(3): doc.add_paragraph()
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run('FRANCHISE DISTRIBUTION SYSTEM'); r.bold=True; r.font.size=Pt(30); r.font.name='Segoe UI'; r.font.color.rgb=RGBColor.from_string('1E3A8A')
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run('Module 1  —  Shop, Pincode, User Limit & Commission'); r.font.size=Pt(15); r.font.name='Segoe UI'; r.font.color.rgb=RGBColor.from_string('2563EB')
doc.add_paragraph()
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run('Complete Flow Documentation'); r.font.size=Pt(13); r.font.color.rgb=RGBColor.from_string('64748B'); r.font.name='Segoe UI'
for _ in range(2): doc.add_paragraph()
pic("d_hier",5.6)
for _ in range(2): doc.add_paragraph()
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run('Prepared for client review  •  read flow-by-flow'); r.italic=True; r.font.size=Pt(10); r.font.color.rgb=RGBColor.from_string('94A3B8')
doc.add_page_break()

# ---------- OVERVIEW ----------
H('Overview — System ek nazar me',17,'1E3A8A')
para([('Ye system ek ',False),('multi-level distribution network',True),(' hai. Sabse upar Warehouse stock deta hai, neeche franchise shops levels me bati hoti hain, har shop kuch pincode serve karti hai, user apne pincode wali shop se khareedta hai, aur har sale pe levels ko commission milta hai. Niche har cheez ',False),('ek-ek karke',True),(' samjhayi gayi hai.',False)])
callout('Structure (BADLA HUA):','Warehouse (WH)  ->  Mini-Warehouse (optional)  ->  Level 1 shop  ->  Level 2 AUR Level 3 (dono seedha Level 1 ke neeche)  ->  User')
H('Is document me kya-kya hai',12,'2563EB',10,2)
for s in ['1. Shop Hierarchy (Warehouse, Mini-WH, Levels)','2. Naming convention + zaroori rules','3. Pincode mapping & rollup','4. User & buying limit','5. Limit increase request (Admin)','6. Commission & money flow (saare cases)','7. Invoice / billing flow']:
    bullet(s)
doc.add_page_break()

# ---------- 1 HIERARCHY ----------
H('1.  Shop Hierarchy',17,'1E3A8A')
para([('Pura network ',False),('tree',True),(' jaisa hai. Top pe ek Warehouse, uske neeche optional Mini-Warehouse, phir Level 1 shops. ',False),('Level 2 aur Level 3 dono seedha Level 1 ke neeche',True),(' hote hain - aapas me bhai-bhai, ek doosre ke neeche nahi. Pehle chain seedhi thi (L1 -> L2 -> L3), ab nahi.',False)])
pic("d_hier",6.4,'Hierarchy: Warehouse -> Mini-WH -> Level 1 (S1,S2,S3) -> S3 ke neeche DONO: S3.1 (Level 2) aur S3.2 (Level 3)')
bullet('Warehouse (WH): sabse upar, saara stock yahi se start hota hai.')
bullet('Mini-Warehouse (MW): optional tier — kuch jagah hota hai. Yeh bhi ek level hai, yahan se bhi khareed sakte, iska apna pincode bhi hota hai. Ek se zyada MW ho sakte.')
bullet('Level 1 shops: franchises. Unke neeche Level 2 aur Level 3 - dono ka parent Level 1 hi hai. Maximum 3 level abhi (future me badha sakte).')
bullet('Levels sparse ho sakti hain — zaroori nahi har Level 1 ke neeche dono (L2 aur L3) hon.')
bullet('Stock lene ka flow same hai — WH, Mini-WH, L1, L2, L3 sab apne upar wale se maal le sakte hain, jaise pehle lete the.')

H('1.1  Multiple Mini-Warehouse',13,'2563EB')
para('Ek se zyada Mini-Warehouse ho sakte. Har MW apni alag shops ko serve karta, stock sabko WH se aata.')
pic("d_multimw",6.2,'WH ke neeche do Mini-Warehouse, har ek apni shops ke saath')

H('1.2  Sparse cases — har shop alag depth tak',13,'2563EB')
para('Zaroori nahi har shop ke teeno level ho. 4 possible cases:')
t=doc.add_table(rows=2,cols=2); t.alignment=WD_ALIGN_PARAGRAPH.CENTER
t.cell(0,0).paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER; t.cell(0,0).paragraphs[0].add_run().add_picture(os.path.join(IMG,'d_spA.png'),width=Inches(2.9))
t.cell(0,1).paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER; t.cell(0,1).paragraphs[0].add_run().add_picture(os.path.join(IMG,'d_spB.png'),width=Inches(2.9))
t.cell(1,0).paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER; t.cell(1,0).paragraphs[0].add_run().add_picture(os.path.join(IMG,'d_spC.png'),width=Inches(2.6))
t.cell(1,1).paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER; t.cell(1,1).paragraphs[0].add_run().add_picture(os.path.join(IMG,'d_spD.png'),width=Inches(2.9))
doc.add_paragraph()
bullet('Case A: sirf Level 1.   Case B: Level 1 + Level 2.   Case C: Level 1 ke neeche dono (L2 aur L3).   Case D: Level 1 + sirf Level 3.')
callout('Purana gap-concept khatam:','Pehle L3 ko L2 ke neeche hona zaroori tha, isliye L2 na hone pe beech me gap banta tha aur uska commission upar roll-up hota tha. Ab L3 seedha L1 ke neeche hai — L2 ka hona ya na hona L3 ko affect hi nahi karta.','FEF3C7','92400E')

H('1.3  Naming convention & rules',13,'2563EB')
para([('Naming: ',False),('S3 (Level 1)',True),('  ->  uske neeche ',False),('S3.1 (Level 2)',True),(' aur ',False),('S3.2 (Level 3)',True),(' — dono ka parent S3 hi hai.',False)])
callout('Rule 1:','Same level pe do shop ka pincode same nahi ho sakta — har pincode level ke andar unique.')
callout('Rule 2:','Har shop apne upper level (parent) se map hoti hai — parent ki shop id ke saath. User ka pincode mapping bhi rehta hai.','FEF3C7','92400E')
doc.add_page_break()

# ---------- 2 PINCODE ----------
H('2.  Pincode Mapping & Rollup',17,'1E3A8A')
para([('Har shop kuch pincode serve karti hai. Neeche wali shop ke pincode ',False),('automatic upar wali shop me bhi map',True),(' ho jaate hain (rollup — neeche se upar).',False)])
pic("d_pin",5.6,'L2 aur L3 dono ka rollup seedha L1 me jaata hai')
table(['Shop','Apne pincode','Neeche se mila','Total serve'],
 [['S3.2 (L3)','p1, p2, p3','—','p1, p2, p3'],
  ['S3.1 (L2)','p4','— (L3 iske neeche nahi)','sirf p4'],
  ['S3 (L1)','p5, p6','p1,p2,p3 (L3 se) + p4 (L2 se)','p1 – p6'],
  ['Mini-WH','p7','p1 – p6 (+S1,S2)','p1 – p9']])
callout('Badla hua:','Pehle L2 ko L3 ke pincode bhi rollup me milte the (kyunki L3 uske neeche tha). Ab nahi — L2 aur L3 alag branch hain, isliye L2 sirf apne pincode serve karta hai. Dono ka rollup seedha L1 me jaata hai.','FEF3C7','92400E')
doc.add_page_break()

# ---------- 3 USER & LIMIT ----------
H('3.  User & Buying Limit',17,'1E3A8A')
para([('User ka ',False),('apna pincode',True),(' hota hai. Rule simple: user ka pincode jis shop ke served-pincode me ho — wahi uski ',False),('home shop',True),(' (unlimited). Pincode alag ho to ',False),('limit',True),(' lagti hai.',False)])
pic("d_user",6.2,'User U2 (pincode p1): green tick wali shops se unlimited, cross wali pe limit')
bullet('U2 ka pincode p1 hai. p1 -> S3.2 (L3), S3 (L1), Mini-WH me match (rollup ki wajah se) -> in sabse UNLIMITED. S3.1 (L2) me nahi — uske paas sirf p4 hai.')
bullet('S1 (p8) aur S2 (p9) me p1 nahi -> in se khareedne pe LIMIT lagti.')

H('3.1  Limit kaise lagti hai',13,'2563EB')
pic("d_limit",4.8,'Har order pe limit check ka flow')
bullet('Home shop (pincode match): koi limit nahi — unlimited.')
bullet('Doosri shop (pincode alag): monthly limit. Default Rs 2000 per month.')
bullet('Effective limit = Admin ne special set kiya to wahi, warna default Rs 2000.')
bullet('Limit rupay (spend) pe count hoti hai, har month 1st ko reset.')
callout('Example:','Limit 2000, ab tak 1500 spend ho chuka, naya order 800 -> 1500+800 = 2300 > 2000 -> BLOCK.','FEE2E2','991B1B')

H('3.2  Limit badhane ki request (Admin)',13,'2563EB')
para('User khud limit nahi badha sakta. Request Admin ke paas jaati hai.')
pic("d_req",5.6,'User request bhejta -> Admin approve/reject')
bullet('User Admin ko request bhejta hai (amount + reason).')
bullet('Approve -> special limit lag jaata (e.g. Rs 10,000 = 2000 normal + 8000 extra), sirf us month ke liye.')
bullet('Next month wapas default Rs 2000. Reject -> koi change nahi.')
doc.add_page_break()

# ---------- 4 COMMISSION ----------
H('4.  Commission & Money Flow',17,'1E3A8A')
callout('Core rule:','Maal har hop pe same rate me transfer hota hai — koi level beech me price nahi badhata. Commission SYSTEM ME FIX NAHI hai: har product ke apne level-rate Admin set karta hai, aur commission usi se nikalta hai. Neeche ke numbers (Rs 800 rate, Rs 100 pool) sirf EXAMPLE hain.')
callout('Structure ka asar:','L2 aur L3 dono seedha L1 ke neeche hain. Isliye ek hi chain me L1+L2+L3 teeno KABHI nahi aate. Har sale me sirf do level: L1+L2, ya L1+L3.','FEF3C7','92400E')
pic("d_setup",5.2,'L1 ke neeche do alag raaste — L2 aur L3')

H('4.1  Raasta 1 — User Level 2 se le',13,'2563EB')
pic("d_cA",6.0)
para('Chain: WH -> L1 -> L2 -> User. Is raaste me L3 aata hi nahi (wo L1 ki doosri branch hai). Sirf do level ko commission: L1 aur L2.')

H('4.2  Raasta 2 — User Level 3 se le',13,'2563EB')
pic("d_cB",6.0)
para('Chain: WH -> L1 -> L3 -> User. Is raaste me L2 aata hi nahi. Yahan bhi sirf do level ko commission: L1 aur L3.')

H('4.3  Case C — Level 1 DIRECT user ko',13,'2563EB')
callout('Rule (sirf Level 1):','L1 jab direct user ko beche -> pura commission 100 ka 50% L1 (=50) aur 50% Warehouse (=50).','FEF3C7','92400E')
pic("d_cC",5.8)

H('4.4  Case D — Level 1 user ko 900 me beche (markup)',13,'2563EB')
para('L1 ne user ko 900 me becha. Base commission 100 -> 50/50 (L1 50, WH 50). Extra 100 (900-800) pura L1 ka. L1 = 150, WH = 50.')
pic("d_cD",6.0)

H('4.5  Case E — Level 3 user ko 900 me beche',13,'2563EB')
para('L3 ne 900 me becha. Base commission normal (L1 + L3 ka apna-apna hissa). Extra 100 pura L3 ka — par report me ye 100 ALAG dikhega (* mark).')
pic("d_cE",6.2)

H('4.6  Commission split - har level ko apna',13,'2563EB')
callout('Rule:','L2 aur L3 ke beech koi rishta nahi hai. Baaki flow bilkul pehle jaisa - har level ko apna set commission milta hai. Sale L2 se -> L1 aur L2 ko apna-apna. Sale L3 se -> L1 aur L3 ko apna-apna. Ek raaste ka level doosre raaste me aata hi nahi, isliye aapas me kuch bantta bhi nahi. Commission ka number product master se aata hai.')
callout('Roll-up KHATAM:','Purana "missing level ka commission upar roll-up ho jata" wala rule ab poori tarah khatam hai. L2 aur L3 alag branch hain - ek ka na hona doosre ko affect hi nahi karta.','FEF3C7','92400E')
pic("d_cF",5.2)

H('4.7  Mini-Warehouse commission',13,'2563EB')
para('Mini-WH chain me ho to uska bhi set commission Rs 15. Maal phir bhi same rate me transfer hota, baaki chain same.')
pic("d_cMW",5.8)

H('4.8  Case G — Mini-WH DIRECT user ko beche',13,'2563EB')
callout('Rule:','Mini-WH agar direct user ko beche -> MW apna 15 rakhta, aur baaki 100 profit WH & MW me 50/50. MW = 15 + 50 = 65, WH = 50.','FEF3C7','92400E')
pic("d_cMWuser",6.0)

H('4.9  Case H — Level 1 direct user (MW chain me)',13,'2563EB')
para('Chain me Mini-WH bhi hai (WH -> MW -> L1), aur L1 direct user ko beche -> MW apna 15, baaki 100 profit WH & L1 me 50/50. MW = 15, WH = 50, L1 = 50.')
pic("d_cL1mw",6.4)

H('4.10  Commission summary',13,'2563EB')
table(['User kahan se khareeda','Kis-kis ko commission','Note'],
 [['Level 2 se','L1 + L2 - dono ko apna-apna','L3 is raaste me aata hi nahi'],
  ['Level 3 se','L1 + L3 - dono ko apna-apna','L2 is raaste me aata hi nahi'],
  ['Level 1 se (direct)','L1 ko apna, baaki 50 Warehouse ko','50/50 rule'],
  ['Mini-WH direct user','MW 65 (15+50), WH 50','50/50 rule'],
  ['L1 direct (MW chain)','L1 50, WH 50, MW 15','MW apna 15 rakhta']])
callout('Dhyan:','L2 aur L3 ab siblings hain, isliye "Level 3 se liya to L2 ko bhi commission" wala purana row ban hi nahi sakta. Roll-up rule bhi khatam. Har level ka rate product master se aata hai - system me fix nahi.','FEF3C7','92400E')
doc.add_page_break()

# ---------- 5 INVOICE ----------
H('5.  Invoice / Billing Flow',17,'1E3A8A')
para('Stock neeche jaate waqt har hop pe document:')
pic("d_inv",6.6)
table(['Hop','Document'],
 [['WH -> Level 1','sirf Stock Transfer note (GST invoice nahi)'],
  ['Level 1 -> Level 2','Tax Invoice (L1 shop naam + GST)'],
  ['Level 2 -> Level 3','Tax Invoice (L2 shop naam + GST)'],
  ['Level 3 -> User','Retail Invoice (L3 shop)']],widths=None)
callout('Rule:','WH -> L1 sirf internal transfer challan. Uske baad har hop pe bechne wali shop ki GST invoice banegi.')

H('Aage kya',13,'2563EB')
para('Yeh Module 1 (Franchise) complete hai. Next: Product module (category, pricing, GST, points) alag document me.')

doc.save(DOCX)
print("SAVED", DOCX)
